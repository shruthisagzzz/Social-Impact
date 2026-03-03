from flask import Flask, request, jsonify, render_template, redirect, url_for, session, flash, send_file
from flask_cors import CORS
import pandas as pd
import numpy as np
import os
import joblib
from transformers import AutoTokenizer, AutoModel
import torch
import sqlite3
from datetime import datetime
import warnings
import google.generativeai as genai
import json
import PyPDF2
import docx
import io
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile

# Suppress warnings
warnings.filterwarnings('ignore')

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # Change this in production
CORS(app)

# Configuration
MODEL_DIR = './models'
DATA_PATH = './data/Final_DataSet.xlsx'
DATABASE = 'social_impact.db'
UPLOAD_FOLDER = './uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

# Configure Gemini API
genai.configure(api_key='AIzaSyBhj6VEH4jGLEdkPtSskE7W1bnr0aCL-YY')
gemini_model = genai.GenerativeModel('gemini-2.0-flash')

# Predefined lists (from your dataset)
COMPANIES = [
    "TCS", "Twitter", "Amazon", "Nestlé", "Accenture", "Coca-Cola", "Facebook",
    "Tata Motors", "Cognizant", "Wipro", "Dell Technologies", "HP", "Flipkart",
    "Biocon", "Cisco", "Tesla", "Samsung", "Intel", "Sony", "Xiaomi", "Google",
    "IBM", "Oracle", "Deloitte", "IKEA", "Tech Mahindra", "Canva", "McAfee",
    "Apple", "Canon", "Razorpay", "PepsiCo", "Salesforce", "Adobe", "Infosys",
    "Microsoft", "Dabur India", "Bosch", "HCL", "Adidas", "Netflix"
]

PROJECT_TYPES = [
    "Education", "Rural Development", "Food Distribution", 
    "Urban Cleanliness", "Health", "Environmental Conservation",
    "Women Empowerment", "Digital Literacy", "Disaster Relief",
    "Renewable Energy", "Wildlife Protection", "Skill Development"
]

# Initialize transformer model
class TextEmbedder:
    def __init__(self, model_name='bert-base-uncased'):
        self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name).to(self.device)
        
    def embed_text(self, text):
        inputs = self.tokenizer(text, return_tensors='pt', truncation=True, 
                              padding=True, max_length=512).to(self.device)
        with torch.no_grad():
            outputs = self.model(**inputs)
        return outputs.last_hidden_state[:,0,:].cpu().numpy()

embedder = TextEmbedder()
models = {}
dataset = None

# Database functions
def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    conn.execute('''
        CREATE TABLE IF NOT EXISTS predictions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            company TEXT NOT NULL,
            project_type TEXT NOT NULL,
            description TEXT,
            sustainability_score REAL,
            community_engagement_score REAL,
            ethical_business_score REAL,
            public_engagement_score REAL,
            total_impact_score REAL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.execute('''
        CREATE TABLE IF NOT EXISTS file_analysis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            filename TEXT NOT NULL,
            domain_type TEXT NOT NULL,
            extracted_text TEXT,
            impact_score REAL,
            grant_recommendation REAL,
            funding_recommendation TEXT,
            gemini_analysis TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.execute('''
        CREATE TABLE IF NOT EXISTS grant_proposals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            project_title TEXT NOT NULL,
            domain_type TEXT NOT NULL,
            requested_amount REAL NOT NULL,
            grant_amount REAL,
            impact_score REAL,
            project_description TEXT,
            executive_summary TEXT,
            objectives TEXT,
            methodology TEXT,
            budget_breakdown TEXT,
            timeline TEXT,
            expected_outcomes TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    conn.commit()
    conn.close()

def load_dataset():
    global dataset
    try:
        if os.path.exists(DATA_PATH):
            if DATA_PATH.endswith('.xlsx'):
                dataset = pd.read_excel(DATA_PATH, engine='openpyxl')
            elif DATA_PATH.endswith('.csv'):
                dataset = pd.read_csv(DATA_PATH)
            print(f"Dataset loaded with {len(dataset)} records")
            return dataset
        else:
            print("Using predefined company list as dataset not found")
            return None
    except Exception as e:
        print(f"Error loading dataset: {str(e)}")
        return None

def load_models():
    global models
    score_names = [
        'Sustainability Score', 
        'Community Engagement Score', 
        'Ethical Business Score', 
        'Public Engagement Score', 
        'Total Impact Score'
    ]
    
    if not os.path.exists(MODEL_DIR):
        os.makedirs(MODEL_DIR)
        print(f"Created models directory: {MODEL_DIR}")
    
    for score_name in score_names:
        filename = score_name.replace(' ', '_').lower() + '_model.pkl'
        model_path = os.path.join(MODEL_DIR, filename)
        
        if os.path.exists(model_path):
            try:
                models[score_name] = joblib.load(model_path)
                print(f"Loaded model for {score_name}")
            except Exception as e:
                print(f"Error loading model for {score_name}: {str(e)}")
                models[score_name] = None
        else:
            print(f"Model not found for {score_name} at {model_path}")
            models[score_name] = None

def predict_scores(company, project_type, description=""):
    try:
        # Generate text embedding if description provided
        if description:
            embedding = embedder.embed_text(description)[0]
        else:
            embedding = np.zeros(768)  # Default zero vector if no description
            
        # Create input features
        input_data = {
            'Company': company,
            'Project Type': project_type,
            **{f'embed_{i}': val for i, val in enumerate(embedding)}
        }
        
        # Convert to DataFrame
        input_df = pd.DataFrame([input_data])
        
        # Make predictions
        predictions = {}
        for score_name, model in models.items():
            if model:
                try:
                    pred = model.predict(input_df)[0]
                    predictions[score_name] = float(np.clip(pred, 0, 100))
                except Exception as e:
                    print(f"Error predicting {score_name}: {str(e)}")
                    predictions[score_name] = 70.0  # Fallback value
            else:
                # Fallback prediction if model not loaded
                base_score = {
                    'Education': 75,
                    'Rural Development': 80,
                    'Food Distribution': 70,
                    'Urban Cleanliness': 65,
                    'Health': 85,
                    'Environmental Conservation': 80,
                    'Women Empowerment': 75,
                    'Digital Literacy': 70,
                    'Disaster Relief': 85,
                    'Renewable Energy': 90,
                    'Wildlife Protection': 85,
                    'Skill Development': 75
                }.get(project_type, 70)
                
                predictions[score_name] = base_score + np.random.uniform(-5, 5)
        
        # Calculate total if not predicted
        if 'Total Impact Score' not in predictions:
            scores = [v for k,v in predictions.items() if k != 'Total Impact Score']
            predictions['Total Impact Score'] = np.mean(scores) if scores else 70.0
            
        return predictions
    except Exception as e:
        print(f"Prediction error: {str(e)}")
        return {
            'Sustainability Score': 70.0,
            'Community Engagement Score': 70.0,
            'Ethical Business Score': 70.0,
            'Public Engagement Score': 70.0,
            'Total Impact Score': 70.0
        }

def save_prediction(user_id, company, project_type, description, scores):
    conn = get_db_connection()
    conn.execute(
        'INSERT INTO predictions (user_id, company, project_type, description, sustainability_score, community_engagement_score, ethical_business_score, public_engagement_score, total_impact_score) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
        (user_id, company, project_type, description, 
         scores['Sustainability Score'], 
         scores['Community Engagement Score'],
         scores['Ethical Business Score'],
         scores['Public Engagement Score'],
         scores['Total Impact Score'])
    )
    conn.commit()
    conn.close()

def get_user_predictions(user_id):
    conn = get_db_connection()
    predictions = conn.execute(
        'SELECT * FROM predictions WHERE user_id = ? ORDER BY created_at DESC',
        (user_id,)
    ).fetchall()
    conn.close()
    return predictions

def get_user_file_analyses(user_id):
    conn = get_db_connection()
    analyses = conn.execute(
        'SELECT * FROM file_analysis WHERE user_id = ? ORDER BY created_at DESC',
        (user_id,)
    ).fetchall()
    conn.close()
    return analyses

def get_user_grant_proposals(user_id):
    conn = get_db_connection()
    proposals = conn.execute(
        'SELECT * FROM grant_proposals WHERE user_id = ? ORDER BY created_at DESC',
        (user_id,)
    ).fetchall()
    conn.close()
    return proposals

def save_file_analysis(user_id, filename, domain_type, extracted_text, impact_score, grant_recommendation, funding_recommendation, gemini_analysis):
    conn = get_db_connection()
    conn.execute(
        'INSERT INTO file_analysis (user_id, filename, domain_type, extracted_text, impact_score, grant_recommendation, funding_recommendation, gemini_analysis) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
        (user_id, filename, domain_type, extracted_text, impact_score, grant_recommendation, funding_recommendation, gemini_analysis)
    )
    conn.commit()
    conn.close()

def save_grant_proposal(user_id, project_title, domain_type, requested_amount, grant_amount, impact_score, project_description, executive_summary, objectives, methodology, budget_breakdown, timeline, expected_outcomes):
    conn = get_db_connection()
    cursor = conn.execute(
        'INSERT INTO grant_proposals (user_id, project_title, domain_type, requested_amount, grant_amount, impact_score, project_description, executive_summary, objectives, methodology, budget_breakdown, timeline, expected_outcomes) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
        (user_id, project_title, domain_type, requested_amount, grant_amount, impact_score, project_description, executive_summary, objectives, methodology, budget_breakdown, timeline, expected_outcomes)
    )
    proposal_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return proposal_id

# File processing functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    filename = file.filename
    if filename.endswith('.pdf'):
        # Extract text from PDF
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif filename.endswith('.docx'):
        # Extract text from DOCX
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    elif filename.endswith('.doc'):
        # For .doc files, we'll try to read as text
        try:
            text = file.read().decode('utf-8')
            return text
        except:
            return "Could not extract text from .doc file. Please upload as .docx or .pdf"
    elif filename.endswith('.txt'):
        # Read text file
        return file.read().decode('utf-8')
    else:
        return "Unsupported file format"

def analyze_with_gemini(domain_type, extracted_text):
    prompt = f"""
    Analyze this social impact project proposal and provide:
    1. An impact score out of 100 (be strict and realistic)
    2. A grant recommendation amount in USD (be reasonable based on the project scope)
    3. Detailed analysis of the project's potential impact, strengths, and areas for improvement
    
    Domain Type: {domain_type}
    Project Proposal:
    {extracted_text[:4000]}  # Limit text to avoid token limits
    
    Please respond in the following JSON format:
    {{
        "impact_score": 85,
        "grant_recommendation": 50000,
        "funding_recommendation": "Breakdown of how the grant should be allocated",
        "analysis": "Detailed analysis of the project..."
    }}
    """
    
    try:
        chat = gemini_model.start_chat(history=[])
        gemini_response = chat.send_message(prompt)
        response_text = gemini_response.text
        
        # Try to extract JSON from the response
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            return json.loads(json_str)
        else:
            # If no JSON found, return default values
            return {
                "impact_score": 70,
                "grant_recommendation": 25000,
                "funding_recommendation": "Standard grant allocation",
                "analysis": response_text
            }
    except Exception as e:
        print(f"Error with Gemini API: {str(e)}")
        return {
            "impact_score": 70,
            "grant_recommendation": 25000,
            "funding_recommendation": "Standard grant allocation",
            "analysis": f"Analysis could not be completed due to an error: {str(e)}"
        }

def generate_grant_proposal_with_gemini(project_title, domain_type, requested_amount, project_description):
    prompt = f"""
    Create a comprehensive grant proposal for a social impact project with the following details:
    
    Project Title: {project_title}
    Domain: {domain_type}
    Requested Amount: ${requested_amount:,.2f}
    Project Description: {project_description}
    
    Please generate a complete grant proposal including:
    1. Executive Summary (concise overview)
    2. Project Objectives (clear, measurable objectives)
    3. Methodology (detailed implementation approach)
    4. Budget Breakdown (itemized budget allocation)
    5. Timeline (project milestones and schedule)
    6. Expected Outcomes (measurable impacts and benefits)
    
    Format the response as JSON:
    {{
        "executive_summary": "Brief overview...",
        "objectives": ["Objective 1", "Objective 2", ...],
        "methodology": "Detailed methodology...",
        "budget_breakdown": "Itemized budget...",
        "timeline": "Project timeline...",
        "expected_outcomes": "Expected outcomes..."
    }}
    """
    
    try:
        chat = gemini_model.start_chat(history=[])
        gemini_response = chat.send_message(prompt)
        response_text = gemini_response.text
        
        # Try to extract JSON from the response
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            return json.loads(json_str)
        else:
            # Fallback proposal if JSON parsing fails
            return {
                "executive_summary": f"This project aims to create significant social impact in the {domain_type} domain through innovative approaches and community engagement.",
                "objectives": [
                    f"Implement effective solutions in {domain_type}",
                    "Engage target communities meaningfully",
                    "Achieve measurable social impact"
                ],
                "methodology": "The project will follow a phased approach including planning, implementation, monitoring, and evaluation stages.",
                "budget_breakdown": f"Total Requested: ${requested_amount:,.2f}\n- Personnel: ${requested_amount*0.4:,.2f}\n- Equipment: ${requested_amount*0.3:,.2f}\n- Operations: ${requested_amount*0.2:,.2f}\n- Contingency: ${requested_amount*0.1:,.2f}",
                "timeline": "Months 1-2: Planning | Months 3-8: Implementation | Months 9-12: Evaluation",
                "expected_outcomes": f"Significant positive impact in {domain_type} with measurable benefits to target communities."
            }
    except Exception as e:
        print(f"Error generating grant proposal: {str(e)}")
        return {
            "executive_summary": f"Project {project_title} aims to address critical needs in {domain_type}.",
            "objectives": ["Deliver meaningful social impact", "Engage stakeholders effectively"],
            "methodology": "Comprehensive project management approach",
            "budget_breakdown": f"Total: ${requested_amount:,.2f}",
            "timeline": "12-month implementation period",
            "expected_outcomes": "Positive community transformation"
        }

def calculate_grant_recommendation(requested_amount, impact_score):
    """Calculate recommended grant amount based on requested amount and impact score"""
    base_multiplier = impact_score / 100
    recommended = requested_amount * base_multiplier
    
    # Apply some reasonable limits
    if recommended > 1000000:  # Cap at 1 million
        recommended = 1000000
    elif recommended < 5000:   # Minimum 5k
        recommended = 5000
        
    return round(recommended, 2)

def create_downloadable_document(proposal_data):
    """Create a professional Word document for the grant proposal"""
    doc = Document()
    
    # Title
    title = doc.add_heading('GRANT PROPOSAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Title
    doc.add_heading(proposal_data['project_title'], level=1)
    
    # Basic Information
    doc.add_heading('Project Information', level=2)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.cell(0, 0).text = 'Domain'
    info_table.cell(0, 1).text = proposal_data['domain_type']
    info_table.cell(1, 0).text = 'Requested Amount'
    info_table.cell(1, 1).text = f"${proposal_data['requested_amount']:,.2f}"
    info_table.cell(2, 0).text = 'Recommended Grant'
    info_table.cell(2, 1).text = f"${proposal_data['grant_amount']:,.2f}"
    info_table.cell(3, 0).text = 'Impact Score'
    info_table.cell(3, 1).text = f"{proposal_data['impact_score']}/100"
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(proposal_data['executive_summary'])
    
    # Project Description
    doc.add_heading('Project Description', level=2)
    doc.add_paragraph(proposal_data['project_description'])
    
    # Objectives
    doc.add_heading('Objectives', level=2)
    for objective in proposal_data['objectives']:
        p = doc.add_paragraph(objective, style='List Bullet')
    
    # Methodology
    doc.add_heading('Methodology', level=2)
    doc.add_paragraph(proposal_data['methodology'])
    
    # Budget Breakdown
    doc.add_heading('Budget Breakdown', level=2)
    doc.add_paragraph(proposal_data['budget_breakdown'])
    
    # Timeline
    doc.add_heading('Project Timeline', level=2)
    doc.add_paragraph(proposal_data['timeline'])
    
    # Expected Outcomes
    doc.add_heading('Expected Outcomes', level=2)
    doc.add_paragraph(proposal_data['expected_outcomes'])
    
    # Save to bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Routes
@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        
        conn = get_db_connection()
        try:
            conn.execute(
                'INSERT INTO users (username, email, password) VALUES (?, ?, ?)',
                (username, email, password)  # In production, hash the password!
            )
            conn.commit()
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            flash('Username or email already exists.', 'error')
        finally:
            conn.close()
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = get_db_connection()
        user = conn.execute(
            'SELECT * FROM users WHERE username = ? AND password = ?',
            (username, password)  # In production, use password hashing!
        ).fetchone()
        conn.close()
        
        if user:
            session['user_id'] = user['id']
            session['username'] = user['username']
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password.', 'error')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    predictions = get_user_predictions(session['user_id'])
    file_analyses = get_user_file_analyses(session['user_id'])
    grant_proposals = get_user_grant_proposals(session['user_id'])
    
    return render_template('dashboard.html', 
                         username=session['username'], 
                         predictions=predictions,
                         file_analyses=file_analyses,
                         grant_proposals=grant_proposals,
                         companies=COMPANIES,
                         project_types=PROJECT_TYPES)

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/predict', methods=['POST'])
def predict():
    if 'user_id' not in session:
        return jsonify({"error": "Authentication required"}), 401
        
    try:
        data = request.json
        company = data.get('company', '').strip()
        project_type = data.get('project_type', '').strip()
        description = data.get('description', '').strip()
        
        if not company or not project_type:
            return jsonify({"error": "Company and project type are required"}), 400
        
        # Get predictions
        scores = predict_scores(company, project_type, description)
        
        # Save prediction to database
        save_prediction(session['user_id'], company, project_type, description, scores)
        
        # Format results
        result = {
            "company": company,
            "project_type": project_type,
            "scores": {k: round(v, 1) for k,v in scores.items()},
            "description_used": bool(description)
        }
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/analyze-file', methods=['POST'])
def analyze_file():
    if 'user_id' not in session:
        return jsonify({"error": "Authentication required"}), 401
        
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
            
        file = request.files['file']
        domain_type = request.form.get('domain_type', '').strip()
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
            
        if not domain_type:
            return jsonify({"error": "Domain type is required"}), 400
            
        if file and allowed_file(file.filename):
            # Extract text from file
            extracted_text = extract_text_from_file(file)
            
            # Analyze with Gemini
            analysis_result = analyze_with_gemini(domain_type, extracted_text)
            
            # Save to database
            save_file_analysis(
                session['user_id'],
                file.filename,
                domain_type,
                extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
                analysis_result['impact_score'],
                analysis_result['grant_recommendation'],
                analysis_result['funding_recommendation'],
                analysis_result['analysis']
            )
            
            return jsonify({
                "filename": file.filename,
                "domain_type": domain_type,
                "impact_score": analysis_result['impact_score'],
                "grant_recommendation": analysis_result['grant_recommendation'],
                "funding_recommendation": analysis_result['funding_recommendation'],
                "analysis": analysis_result['analysis']
            })
        else:
            return jsonify({"error": "Invalid file type. Allowed: txt, pdf, doc, docx"}), 400
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/generate-grant-proposal', methods=['POST'])
def generate_grant_proposal():
    if 'user_id' not in session:
        return jsonify({"error": "Authentication required"}), 401
        
    try:
        data = request.json
        project_title = data.get('project_title', '').strip()
        domain_type = data.get('domain_type', '').strip()
        requested_amount = float(data.get('requested_amount', 0))
        project_description = data.get('project_description', '').strip()
        
        if not project_title or not domain_type or requested_amount <= 0:
            return jsonify({"error": "Project title, domain type, and valid requested amount are required"}), 400
        
        # Calculate impact score (you can modify this logic)
        impact_score = min(100, max(20, (requested_amount / 10000) * 10 + 50))  # Example calculation
        
        # Generate grant proposal with Gemini
        proposal_data = generate_grant_proposal_with_gemini(
            project_title, domain_type, requested_amount, project_description
        )
        
        # Calculate recommended grant amount
        grant_amount = calculate_grant_recommendation(requested_amount, impact_score)
        
        # Convert objectives to string if it's a list
        objectives = proposal_data['objectives']
        if isinstance(objectives, list):
            objectives = '\n'.join(objectives)
        
        # Ensure all fields are strings
        executive_summary = str(proposal_data['executive_summary'])
        methodology = str(proposal_data['methodology'])
        budget_breakdown = str(proposal_data['budget_breakdown'])
        timeline = str(proposal_data['timeline'])
        expected_outcomes = str(proposal_data['expected_outcomes'])
        
        # Save to database
        proposal_id = save_grant_proposal(
            session['user_id'],
            project_title,
            domain_type,
            requested_amount,
            grant_amount,
            impact_score,
            project_description,
            executive_summary,
            objectives,
            methodology,
            budget_breakdown,
            timeline,
            expected_outcomes
        )
        
        return jsonify({
            "proposal_id": proposal_id,
            "project_title": project_title,
            "domain_type": domain_type,
            "requested_amount": requested_amount,
            "grant_amount": grant_amount,
            "impact_score": impact_score,
            "executive_summary": executive_summary,
            "objectives": objectives,
            "methodology": methodology,
            "budget_breakdown": budget_breakdown,
            "timeline": timeline,
            "expected_outcomes": expected_outcomes
        })
        
    except Exception as e:
        print(f"Error generating grant proposal: {str(e)}")
        return jsonify({"error": str(e)}), 500

def save_grant_proposal(user_id, project_title, domain_type, requested_amount, grant_amount, impact_score, project_description, executive_summary, objectives, methodology, budget_breakdown, timeline, expected_outcomes):
    conn = get_db_connection()
    cursor = conn.execute(
        '''INSERT INTO grant_proposals 
        (user_id, project_title, domain_type, requested_amount, grant_amount, impact_score, 
         project_description, executive_summary, objectives, methodology, budget_breakdown, 
         timeline, expected_outcomes) 
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
        (user_id, project_title, domain_type, requested_amount, grant_amount, impact_score,
         str(project_description), str(executive_summary), str(objectives), str(methodology),
         str(budget_breakdown), str(timeline), str(expected_outcomes))
    )
    proposal_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return proposal_id
@app.route('/download-proposal/<int:proposal_id>')
def download_proposal(proposal_id):
    if 'user_id' not in session:
        return jsonify({"error": "Authentication required"}), 401
        
    try:
        conn = get_db_connection()
        proposal = conn.execute(
            'SELECT * FROM grant_proposals WHERE id = ? AND user_id = ?',
            (proposal_id, session['user_id'])
        ).fetchone()
        conn.close()
        
        if not proposal:
            return jsonify({"error": "Proposal not found"}), 404
        
        # Prepare data for document generation
        proposal_data = {
            'project_title': proposal['project_title'],
            'domain_type': proposal['domain_type'],
            'requested_amount': proposal['requested_amount'],
            'grant_amount': proposal['grant_amount'],
            'impact_score': proposal['impact_score'],
            'project_description': proposal['project_description'],
            'executive_summary': proposal['executive_summary'],
            'objectives': proposal['objectives'].split('\n') if proposal['objectives'] else [],
            'methodology': proposal['methodology'],
            'budget_breakdown': proposal['budget_breakdown'],
            'timeline': proposal['timeline'],
            'expected_outcomes': proposal['expected_outcomes']
        }
        
        # Create downloadable document
        document_buffer = create_downloadable_document(proposal_data)
        
        filename = f"grant_proposal_{proposal['project_title'].replace(' ', '_')}.docx"
        
        return send_file(
            document_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/companies', methods=['GET'])
def get_companies():
    global dataset
    try:
        if dataset is not None:
            for col in ['Company', 'company', 'COMPANY']:
                if col in dataset.columns:
                    companies = dataset[col].dropna().unique().tolist()
                    return jsonify(sorted(list(set(companies + COMPANIES))))
        return jsonify(sorted(COMPANIES))
    except Exception as e:
        print(f"Error getting companies: {str(e)}")
        return jsonify(sorted(COMPANIES))

@app.route('/api/project-types', methods=['GET'])
def get_project_types():
    global dataset
    try:
        if dataset is not None:
            for col in ['Project Type', 'project_type', 'PROJECT_TYPE']:
                if col in dataset.columns:
                    types = dataset[col].dropna().unique().tolist()
                    return jsonify(sorted(list(set(types + PROJECT_TYPES))))
        return jsonify(sorted(PROJECT_TYPES))
    except Exception as e:
        print(f"Error getting project types: {str(e)}")
        return jsonify(sorted(PROJECT_TYPES))

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "models_loaded": len([m for m in models.values() if m is not None]),
        "dataset_loaded": dataset is not None,
        "companies_count": len(COMPANIES),
        "project_types_count": len(PROJECT_TYPES)
    })

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

if __name__ == '__main__':
    # Create upload directory
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    
    # Initialize database
    init_db()
    
    # Load dataset and models
    print("Loading dataset...")
    load_dataset()
    
    print("Loading models...")
    load_models()
    
    # Start server
    print("\nAPI Ready with:")
    print(f"- {len(COMPANIES)} companies")
    print(f"- {len(PROJECT_TYPES)} project types")
    print(f"- {len([m for m in models.values() if m is not None])} models loaded")
    print("\nStarting server on http://localhost:5000")
    app.run(host='0.0.0.0', port=5000, debug=True)